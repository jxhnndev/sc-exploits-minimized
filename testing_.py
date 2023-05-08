import os
import PyPDF2
from docx import Document
from pyth.plugins.rtf15.reader import Rtf15Reader
from pyth.plugins.plaintext.writer import PlaintextWriter


def search_text_in_file(file_path, search_text):
    """Ищет заданный текст в файле"""
    file_extension = file_path.split('.')[-1].lower()

    # Проверяем расширение файла и вызываем соответствующую функцию для чтения текста из файла
    if file_extension == 'pdf':
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            num_pages = pdf_reader.getNumPages()
            for page_num in range(num_pages):
                page = pdf_reader.getPage(page_num)
                text = page.extractText()
                if search_text in text:
                    return True
    elif file_extension == 'docx':
        document = Document(file_path)
        for para in document.paragraphs:
            if search_text in para.text:
                return True
    elif file_extension == 'doc':
        document = Document(file_path)
        for para in document.paragraphs:
            if search_text in para.text:
                return True
    elif file_extension == 'rtf':
        with open(file_path, 'rb') as rtf_file:
            doc = Rtf15Reader.read(rtf_file)
            text = PlaintextWriter.write(doc).getvalue()
            if search_text in text:
                return True
    elif file_extension == 'odt':
        with open(file_path, 'rb') as odt_file:
            doc = Rtf15Reader.read(odt_file)
            text = PlaintextWriter.write(doc).getvalue()
            if search_text in text:
                return True

    return False


def search_text_in_folder(folder_path, search_text):
    """Ищет заданный текст во всех файлах заданных форматов в директории"""
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            if file_path.lower().endswith(('.pdf', '.doc', '.docx', '.rtf', '.odt')):
                if search_text_in_file(file_path, search_text):
                    print(f'Текст найден в файле {file_path}')


search_text_in_folder('/Users/caramba/Desktop/ALL_DATA', 'решение')
